from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
import pdfplumber
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import io
import re
import os
from dotenv import load_dotenv

from database import (
    init_db, list_profiles, create_profile,
    get_profile, update_profile_data, delete_profile_by_id,
    list_applications, create_application, update_application,
    delete_application, get_application_stats,
    list_events, create_event, delete_event,
)

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

API_KEY = os.getenv("SILICONFLOW_API_KEY")
BASE_URL = "https://api.siliconflow.cn/v1"
MODEL = "deepseek-ai/DeepSeek-V3"

client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

init_db()


# ─── 工具函数 ────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        return file_bytes.decode("utf-8", errors="ignore")


def desc_to_bullets(text: str) -> list:
    if not text:
        return []
    lines = re.split(r'\n|(?<=[.。])\s+(?=[A-Z\u4e00-\u9fa5])', text.strip())
    bullets = [l.strip().lstrip('-•·').strip() for l in lines if len(l.strip()) > 5]
    return bullets if bullets else [text.strip()]


def normalize_profile(data: dict) -> dict:
    for item in data.get("experience", []):
        if "bullets" not in item or not item["bullets"]:
            item["bullets"] = desc_to_bullets(item.pop("description", ""))
        item.setdefault("location", "无")
    for item in data.get("projects", []):
        if "bullets" not in item or not item["bullets"]:
            item["bullets"] = desc_to_bullets(item.pop("description", ""))
        item.setdefault("start", "无")
        item.setdefault("end", "无")
        item.setdefault("location", "无")
    for item in data.get("activities", []):
        if "bullets" not in item or not item["bullets"]:
            item["bullets"] = desc_to_bullets(item.pop("description", ""))
        item.setdefault("location", "无")
        item.setdefault("start", "无")
        item.setdefault("end", "无")
    for item in data.get("education", []):
        item.setdefault("location", "无")
    return data


# ─── 简历解析（只解析，不保存）────────────────────────────

@app.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    file_bytes = await file.read()
    resume_text = extract_text(file_bytes, file.filename)

    prompt = f"""你是简历解析器，只输出JSON，不要有任何其他文字。

解析以下简历，严格按照格式输出，所有字段找不到填"无"：

{{
  "basic": {{"name":"","email":"","phone":"","location":"","linkedin":""}},
  "education": [{{"school":"","location":"学校所在城市","degree":"","major":"","gpa":"","start":"","end":"","courses":""}}],
  "experience": [{{"company":"","location":"公司所在城市","title":"","type":"实习或全职","start":"","end":"","bullets":["职责1","职责2"]}}],
  "projects": [{{"name":"","location":"无","tech":"","start":"","end":"","bullets":["描述1","描述2"]}}],
  "activities": [{{"organization":"","role":"","location":"","start":"","end":"","bullets":["描述1"]}}],
  "skills": {{"technical":"技能1,技能2","languages":"语言能力"}},
  "awards": [{{"name":"","date":""}}]
}}

简历内容：
{resume_text}"""

    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        return {"status": "error", "message": f"JSON解析失败: {e}", "raw": raw[:300]}

    data = normalize_profile(data)
    return {"status": "ok", "data": data}


# ─── 简历优化 ────────────────────────────────────────────

@app.post("/analyze")
async def analyze(file: UploadFile = File(...), jd: str = Form(...)):
    file_bytes = await file.read()
    resume_text = extract_text(file_bytes, file.filename)

    prompt = f"""你是一位专业的求职顾问，帮助应届生优化简历。

【我的简历】
{resume_text}

【目标岗位JD】
{jd}

请按以下格式输出分析结果：

## 命中关键词
列出简历中已经匹配JD要求的技能/经历（每条一行，用✅开头）

## 优化建议
针对JD要求，指出简历中可以强化表达的地方（每条一行，用📝开头）

## 缺口分析
列出JD要求但简历中明显缺失的内容（每条一行，用⚠️开头）

## 优化后的简历摘要
根据JD重新撰写一段个人简介（100-150字），突出与该岗位最相关的经历和能力。
"""

    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
    )

    return {"result": response.choices[0].message.content}


# ─── 简历优化生成 ────────────────────────────────────────

@app.post("/optimize-resume")
async def optimize_resume(file: UploadFile = File(...), jd: str = Form(...)):
    file_bytes = await file.read()
    resume_text = extract_text(file_bytes, file.filename)

    prompt = f"""你是一位专业的简历改写专家。请根据目标岗位JD，将原始简历内容针对性地改写优化，使其更匹配该岗位。

要求：
- 保留所有真实经历，不虚构内容
- 在每条工作/项目描述中融入JD的核心关键词
- 用更有力的动词开头（如：主导、设计、构建、优化、协作）
- 量化成果（如有数据则保留，无数据则合理强化描述）
- 只输出JSON，不要有任何其他文字

严格按照以下格式输出（所有字段找不到填"无"）：

{{
  "basic": {{"name":"","email":"","phone":"","location":"","linkedin":""}},
  "education": [{{"school":"","location":"","degree":"","major":"","gpa":"","start":"","end":"","courses":""}}],
  "experience": [{{"company":"","location":"","title":"","type":"","start":"","end":"","bullets":["改写后的职责1","职责2"]}}],
  "projects": [{{"name":"","location":"无","tech":"","start":"","end":"","bullets":["改写后的描述1","描述2"]}}],
  "activities": [{{"organization":"","role":"","location":"","start":"","end":"","bullets":["描述1"]}}],
  "skills": {{"technical":"技能1,技能2","languages":"语言能力"}},
  "awards": [{{"name":"","date":""}}]
}}

【原始简历】
{resume_text}

【目标岗位JD】
{jd}"""

    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        return {"status": "error", "message": f"JSON解析失败: {e}", "raw": raw[:300]}

    data = normalize_profile(data)
    return {"status": "ok", "data": data}


def _add_horizontal_line(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    _add_horizontal_line(doc)


def _bullet_para(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.size = Pt(9.5)


@app.post("/build-docx")
async def build_docx(payload: dict):
    data = payload.get("data", {})
    basic = data.get("basic", {})

    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ── 姓名 ──
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(basic.get("name", "姓名"))
    name_run.bold = True
    name_run.font.size = Pt(18)

    # ── 联系信息 ──
    contact_parts = [x for x in [
        basic.get("email"), basic.get("phone"),
        basic.get("location"), basic.get("linkedin")
    ] if x and x != "无"]
    contact_p = doc.add_paragraph(" | ".join(contact_parts))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(6)
    for run in contact_p.runs:
        run.font.size = Pt(9)

    # ── 教育经历 ──
    edu_list = [e for e in data.get("education", []) if e.get("school") and e["school"] != "无"]
    if edu_list:
        _section_heading(doc, "教育经历")
        for edu in edu_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(edu.get("school", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            date_str = f"{edu.get('start','')} – {edu.get('end','')}"
            tab_run = p.add_run(f"\t{date_str}")
            tab_run.font.size = Pt(9.5)
            p.add_run(f"\n{edu.get('degree','')} · {edu.get('major','')}").font.size = Pt(9.5)
            if edu.get("gpa") and edu["gpa"] != "无":
                p.add_run(f"   GPA: {edu['gpa']}").font.size = Pt(9.5)
            if edu.get("courses") and edu["courses"] != "无":
                _bullet_para(doc, f"主要课程：{edu['courses']}")

    # ── 工作/实习经历 ──
    exp_list = [e for e in data.get("experience", []) if e.get("company") and e["company"] != "无"]
    if exp_list:
        _section_heading(doc, "工作经历")
        for exp in exp_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(exp.get("company", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            date_str = f"{exp.get('start','')} – {exp.get('end','')}"
            p.add_run(f"\t{date_str}").font.size = Pt(9.5)
            title_type = exp.get("title", "")
            if exp.get("type") and exp["type"] != "无":
                title_type += f" · {exp['type']}"
            p.add_run(f"\n{title_type}").font.size = Pt(9.5)
            for b in exp.get("bullets", []):
                if b and b != "无":
                    _bullet_para(doc, b)

    # ── 项目经历 ──
    proj_list = [p for p in data.get("projects", []) if p.get("name") and p["name"] != "无"]
    if proj_list:
        _section_heading(doc, "项目经历")
        for proj in proj_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            date_str = f"{proj.get('start','')} – {proj.get('end','')}"
            p.add_run(f"\t{date_str}").font.size = Pt(9.5)
            if proj.get("tech") and proj["tech"] != "无":
                p.add_run(f"\n技术栈：{proj['tech']}").font.size = Pt(9.5)
            for b in proj.get("bullets", []):
                if b and b != "无":
                    _bullet_para(doc, b)

    # ── 课外活动 ──
    act_list = [a for a in data.get("activities", []) if a.get("organization") and a["organization"] != "无"]
    if act_list:
        _section_heading(doc, "校园活动")
        for act in act_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(act.get("organization", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            date_str = f"{act.get('start','')} – {act.get('end','')}"
            p.add_run(f"\t{date_str}").font.size = Pt(9.5)
            p.add_run(f"\n{act.get('role','')}").font.size = Pt(9.5)
            for b in act.get("bullets", []):
                if b and b != "无":
                    _bullet_para(doc, b)

    # ── 技能 ──
    skills = data.get("skills", {})
    if skills.get("technical") and skills["technical"] != "无":
        _section_heading(doc, "技能")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.add_run("技术技能：").bold = True
        p.add_run(skills["technical"]).font.size = Pt(9.5)
        if skills.get("languages") and skills["languages"] != "无":
            p2 = doc.add_paragraph()
            p2.add_run("语言能力：").bold = True
            p2.add_run(skills["languages"]).font.size = Pt(9.5)

    # ── 荣誉奖项 ──
    award_list = [a for a in data.get("awards", []) if a.get("name") and a["name"] != "无"]
    if award_list:
        _section_heading(doc, "荣誉奖项")
        for award in award_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            r1 = p.add_run(award.get("name", ""))
            r1.font.size = Pt(9.5)
            if award.get("date") and award["date"] != "无":
                p.add_run(f"\t{award['date']}").font.size = Pt(9.5)

    # 输出为字节流
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    name = basic.get("name", "resume") or "resume"
    filename = f"{name}_optimized.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )


# ─── 档案 CRUD ───────────────────────────────────────────

@app.get("/profiles")
async def api_list_profiles():
    return list_profiles()


@app.post("/profiles")
async def api_create_profile(payload: dict):
    name = payload.get("name", "未命名档案")
    data = payload.get("data", {})
    profile_id = create_profile(name, data)
    return {"id": profile_id, "name": name}


@app.get("/profiles/{profile_id}")
async def api_get_profile(profile_id: int):
    data = get_profile(profile_id)
    if not data:
        return JSONResponse(status_code=404, content={"error": "档案不存在"})
    return data


@app.put("/profiles/{profile_id}")
async def api_update_profile(profile_id: int, payload: dict):
    update_profile_data(profile_id, payload)
    return {"status": "ok"}


@app.delete("/profiles/{profile_id}")
async def api_delete_profile(profile_id: int):
    delete_profile_by_id(profile_id)
    return {"status": "ok"}


# ─── 投递记录 CRUD ───────────────────────────────────────

@app.get("/applications")
async def api_list_applications():
    return list_applications()

@app.post("/applications")
async def api_create_application(payload: dict):
    app_id = create_application(payload)
    return {"id": app_id}

@app.put("/applications/{app_id}")
async def api_update_application(app_id: int, payload: dict):
    update_application(app_id, payload)
    return {"status": "ok"}

@app.delete("/applications/{app_id}")
async def api_delete_application(app_id: int):
    delete_application(app_id)
    return {"status": "ok"}

@app.get("/applications/stats")
async def api_application_stats():
    return get_application_stats()

@app.get("/applications/{app_id}/events")
async def api_list_events(app_id: int):
    return list_events(app_id)

@app.post("/applications/{app_id}/events")
async def api_create_event(app_id: int, payload: dict):
    event_id = create_event(app_id, payload)
    return {"id": event_id}

@app.delete("/applications/{app_id}/events/{event_id}")
async def api_delete_event(app_id: int, event_id: int):
    delete_event(event_id)
    return {"status": "ok"}


# ─── 页面路由 ────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def index():
    with open("index.html", "r", encoding="utf-8") as f:
        return f.read()


@app.get("/profile-page", response_class=HTMLResponse)
async def profile_page():
    with open("profile.html", "r", encoding="utf-8") as f:
        return f.read()


@app.get("/tracker", response_class=HTMLResponse)
async def tracker_page():
    with open("tracker.html", "r", encoding="utf-8") as f:
        return f.read()


@app.get("/test-form", response_class=HTMLResponse)
async def test_form():
    with open("test_form.html", "r", encoding="utf-8") as f:
        return f.read()
